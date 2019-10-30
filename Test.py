from zipfile import ZipFile
from bs4 import BeautifulSoup
import copy
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from win32com import client

#批量将目录下的所有doc文档转化成docx文档
def getfile1():
    for root, dirs, files in os.walk(os.getcwd()):
        for file in files:
            if file[-4:] == '.doc' and '~$' not in file:
                print('正在把doc文件转换成docx文件 {}'.format(os.path.join(root, file)))
                word = client.Dispatch('Word.Application')
                doc = word.Documents.Open(os.path.join(root, file))
                doc.SaveAs(os.path.join(root, file[:-4]), 12, False, "", True, "", False, False, False,
                           False)
                # 转化后路径下的文件
                doc.Close()
getfile1()

#批量删除原来的doc文档
def getfile2():
    for root, dirs, files in os.walk(os.getcwd()):
        for name in files:
            if name.endswith(".doc"):
                os.remove(os.path.join(root, name))
                print("正在删除doc文件: " + os.path.join(root, name))
getfile2()


# def main():
path = input('请输入待处理文件路径:');
# 原文件名称
filepaths = [];
os.listdir(path)
for file in os.listdir(path):
    file_path = os.path.join(path, file)
    # if not os.path.isdir(file_path) and os.path.splitext(file_path)[1] == '.doc':
    #     doc_to_docx(file_path);

    if not os.path.isdir(file_path) and os.path.splitext(file_path)[1] == '.docx' \
       and file[0:2] != '~$' and file[0:2] != '.~':
        filepaths.append(file_path);
    # print(oldnames)

if not os.path.exists(path + 'result/'):
    os.mkdir(path + 'result/');
for filepath in filepaths:

    document = ZipFile(filepath)
    xml = document.read("word/document.xml")
    wordObj = BeautifulSoup(xml.decode("utf-8"))
    # 插入：w:ins，删除：w:del，移动 ：w:moveFrom，w:moveTo，设置格式：w:rFonts，批注：w:commentRangeEnd
    paras = wordObj.find_all("w:p");

    oldfilename = filepath.split('/')[-1];
    first = oldfilename.split('_')[0]
    firstA = oldfilename.split('_')[0] + '_修改前后对照表.docx'


    # 存入的文档
    doc = docx.Document();
    # 标题
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('《'+first+'》'+'\n'+'修订前后对照表')
    run.font.name = '宋体';
    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体');
    run.font.size = docx.shared.Pt(14)

    # 查找标题id
    stylexml = document.read("word/styles.xml")
    styleObj = BeautifulSoup(stylexml.decode("utf-8"))
    headingstyle =styleObj.find(attrs={'w:val' : 'heading 1'})
    headingstyleid = headingstyle.parent.get('w:styleid')


    # 表格
    table = doc.add_table(rows=1, cols=3, style=None);
    table.border = docx.shared.Pt(0.5);
    table.style = 'TableGrid';
    doc.styles['Normal'].font.name = '宋体';
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体');
    doc.styles['Normal'].font.size = docx.shared.Pt(10.5)
    p0 = table.cell(0, 0).add_paragraph();
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER;
    p0run = p0.add_run('章节');
    p0run.font.bold=True;
    p1 = table.cell(0, 1).add_paragraph();
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER;
    p1run = p1.add_run('修订前内容');
    p1run.font.bold=True;
    p2 = table.cell(0, 2).add_paragraph();
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER;
    p2run = p2.add_run('修订后内容');
    p2run.font.bold=True;

    # 章节名称
    charpter = '';
    # 章节开始行
    charpterstartindex = 1;
    for para in paras:

        # 查找插入
        wins = para.find_all("w:ins");
        # 查找删除
        wdel = para.find_all("w:del");
        if (len(wins) and wins[0].get_text() != "") > 0 or (len(wdel) > 0 and wdel[0].get_text() != ""):
            # 查找前面第一个标题
            wpparas = para.find_all_previous('w:p');
            if len(wpparas)>0:
                haschapter = False
                for index in range(len(wpparas)):
                    wppara = wpparas[index]
                    p = wppara.find('w:pstyle',attrs={'w:val' : headingstyleid})
                    if p is not None :
                        if wppara.get_text() is not None and charpter !='' and  charpter != wppara.get_text():
                             # 合并单元格
                            table.cell(charpterstartindex, 0).merge(table.cell(len(table.rows)-1, 0)).text = charpter
                            charpterstartindex = len(table.rows)
                        charpter = wppara.get_text();
                        haschapter = True
                        break
                    if index == (len(wpparas)-1) and not haschapter:
                        charpterstartindex += 1
            oldpar = copy.copy(para)
            newpar = copy.copy(para)
            # 原文内容：删掉添加的
            oldcontents = oldpar.find_all("w:ins")
            # 修订版内容：去掉删除的
            newcontents = newpar.find_all("w:del")
            for oldcontent in oldcontents:
                oldcontent.extract()
            for newcontent in newcontents:
                newcontent.extract()
            row = table.add_row()
            row.cells[0].text = charpter
            row.cells[1].text = oldpar.get_text()
            row.cells[2].text = newpar.get_text()


    doc.save(path + 'result/' + firstA);
    # if __name__ == "main()":
    #     main();
